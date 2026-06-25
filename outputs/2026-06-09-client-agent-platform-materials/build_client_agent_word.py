from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCES = [
    ("第一部分：平台补充页", ROOT / "sage-agent-platform-addendum-zh.md"),
    ("第二部分：沟通弹药库", ROOT / "sage-agent-platform-battlecard-zh.md"),
    ("第三部分：框架差异性", ROOT / "sage-framework-differentiation-zh.md"),
]
OUT = ROOT / "面向某某的Sage-Agent平台沟通材料.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x11, 0x1B, 0x2B)
MUTED = RGBColor(0x5A, 0x64, 0x73)
FILL = "E8EEF5"
CALLOUT = "F4F6F9"
BORDER = "CAD6E2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)


def set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_run_font(run, size=None, bold=None, color=None) -> None:
    run.font.name = "Calibri"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")


def add_runs_with_inline_code(
    paragraph, text: str, size=11, color=INK, bold=False
) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        is_code = part.startswith("`") and part.endswith("`")
        run = paragraph.add_run(part[1:-1] if is_code else part)
        set_run_font(run, size=size, color=color, bold=bold)
        if is_code:
            run.font.name = "Consolas"
            run.font.color.rgb = DARK_BLUE
            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)
            r_fonts.set(qn("w:eastAsia"), "Consolas")
            r_fonts.set(qn("w:ascii"), "Consolas")
            r_fonts.set(qn("w:hAnsi"), "Consolas")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = INK
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.95)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("Sage Agent 平台沟通材料")
    set_run_font(run, 28, True, RGBColor(0x0B, 0x25, 0x45))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("面向 某某 银行高管与 IT 团队的理解弹药库")
    set_run_font(run, 15, False, DARK_BLUE)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=180, bottom=180, start=220, end=220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("核心定位")
    set_run_font(run, 12, True, DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_runs_with_inline_code(
        p,
        "Sage 不只是把大模型包装成聊天或工具调用，而是把 Agent 定义成一个可产品化、可治理、可复用、可观察的任务执行单元。",
        size=11,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("文档结构")
    set_run_font(run, 13, True, BLUE)

    for item in [
        "平台补充页：可直接抽取到 PPT 的新增页文案",
        "沟通弹药库：面向高管、销售、方案和会后邮件的论点素材",
        "框架差异性：Sage 与普通 Agent 框架、RAG、RPA、单点 Copilot 的关键差异",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_runs_with_inline_code(p, item)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("版本：2026-06-10 | 内部沟通草案")
    set_run_font(run, 9, False, MUTED)

    doc.add_page_break()


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.text = ""
        run = p.add_run("Sage Agent Platform | 某某 Discussion Materials")
        set_run_font(run, 8, False, MUTED)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = ""
        run = fp.add_run("Internal discussion draft")
        set_run_font(run, 8, False, MUTED)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    rows = []
    i = start
    while (
        i < len(lines)
        and lines[i].strip().startswith("|")
        and lines[i].strip().endswith("|")
    ):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    if len(rows) < 2:
        return None
    sep = rows[1]
    if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in sep):
        return None
    return [rows[0]] + rows[2:], i


def column_widths(headers: list[str], col_count: int) -> list[int]:
    if col_count == 2:
        return [2700, 6660]
    if col_count == 3:
        return [2160, 3600, 3600]
    if col_count == 4:
        return [1600, 2580, 2580, 2600]
    return [int(9360 / col_count)] * col_count


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    if len(rows) >= 8:
        doc.add_page_break()
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    widths = column_widths(rows[0], col_count)
    set_table_width(table, widths)

    for r_idx, row in enumerate(rows):
        set_row_cant_split(table.rows[r_idx])
        if r_idx == 0:
            set_row_repeat_header(table.rows[r_idx])
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_width(cell, widths[c_idx])
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if r_idx == 0:
                set_cell_shading(cell, FILL)
                add_runs_with_inline_code(p, text, size=10, color=DARK_BLUE, bold=True)
            else:
                add_runs_with_inline_code(p, text, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_runs_with_inline_code(p, text, size=10.5, color=DARK_BLUE, bold=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_markdown_file(doc: Document, title: str, path: Path, first=False) -> None:
    if not first:
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        add_header_footer(doc)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.add_run(title)

    lines = path.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if line.startswith("```"):
            if in_code:
                add_callout(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not line:
            i += 1
            continue

        parsed = parse_table(lines, i)
        if parsed:
            rows, i = parsed
            add_markdown_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1:
                style = "Heading 2"
            elif level == 2:
                style = "Heading 2"
            else:
                style = "Heading 3"
            p = doc.add_paragraph(style=style)
            add_runs_with_inline_code(
                p,
                text,
                size=13 if style == "Heading 2" else 12,
                color=BLUE if style == "Heading 2" else DARK_BLUE,
                bold=True,
            )
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_inline_code(p, line[2:].strip())
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_inline_code(p, numbered.group(1).strip())
            i += 1
            continue

        if line.startswith("> "):
            add_callout(doc, line[2:].strip())
            i += 1
            continue

        if line.startswith("`") and line.endswith("`"):
            add_callout(doc, line.strip("`"))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_inline_code(p, line)
        i += 1


def main() -> None:
    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    # First real section starts after the cover page.
    add_header_footer(doc)
    for idx, (title, path) in enumerate(SOURCES):
        add_markdown_file(doc, title, path, first=(idx == 0))

    # Final preset audit note for maintainers.
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_header_footer(doc)
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("附录：排版说明")
    for item in [
        "文档采用 compact_reference_guide 风格：适合密集型销售材料、沟通手册和备讲指南。",
        "正文使用 11 pt 字号、1.25 倍行距；标题使用蓝色层级，便于快速扫描。",
        "Markdown 表格已转换为真实 Word 表格，并设置固定宽度、表头底色和单元格内边距。",
        "代码式流程和关键表达被处理为浅色提示框，便于在会前快速定位。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_runs_with_inline_code(p, item)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
