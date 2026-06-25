"""
DOCX文件解析器
支持Word文档的文本提取和元数据获取
"""

import traceback
from typing import Dict, Any
from docx import Document
from .base_parser import BaseFileParser, ParseResult


class DOCXParser(BaseFileParser):
    """DOCX文件解析器"""

    SUPPORTED_EXTENSIONS = [".docx"]
    SUPPORTED_MIME_TYPES = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    def parse(self, file_path: str, skip_validation: bool = False) -> ParseResult:
        """
        解析DOCX文件

        Args:
            file_path: DOCX文件路径
            skip_validation: 是否跳过文件格式验证（can_parse检查）

        Returns:
            ParseResult: 解析结果
        """
        if not self.validate_file(file_path):
            return self.create_error_result(
                f"文件不存在或无法读取: {file_path}", file_path
            )

        if not skip_validation and not self.can_parse(file_path):
            return self.create_error_result(f"不支持的文件类型: {file_path}", file_path)

        try:
            doc = Document(file_path)

            # 提取文本内容
            text = self._extract_text(doc)

            # 提取DOCX元数据
            docx_metadata = self._extract_docx_metadata(doc)

            # 合并基础元数据和DOCX特定元数据
            base_metadata = self.get_file_metadata(file_path)
            metadata = {**base_metadata, **docx_metadata}

            # 添加文本统计信息
            metadata.update(
                {
                    "text_length": len(text),
                    "character_count": len(text),
                    "word_count": len(text.split()) if text else 0,
                    "line_count": text.count("\\n") if text else 0,
                }
            )

            return ParseResult(text=text, metadata=metadata, success=True)

        except Exception as e:
            error_msg = f"DOCX解析失败: {str(e)}"
            print(error_msg)
            traceback.print_exc()
            return self.create_error_result(error_msg, file_path)

    def _extract_text(self, doc: Any) -> str:
        """
        提取文档文本内容

        Args:
            doc: Word文档对象

        Returns:
            str: 提取的文本
        """
        text_parts = []

        try:
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 提取表格文本
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(f"\\n--- 表格 ---\\n{table_text}")

        except Exception as e:
            print(f"提取DOCX文本时出错: {e}")
            traceback.print_exc()

        return "\\n".join(text_parts)

    def _extract_table_text(self, table) -> str:
        """
        提取表格文本

        Args:
            table: 表格对象

        Returns:
            str: 表格文本
        """
        table_text = []

        try:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_text.append(cell_text)
                if any(row_text):  # 只添加非空行
                    table_text.append(" | ".join(row_text))
        except Exception as e:
            print(f"提取表格文本时出错: {e}")
            traceback.print_exc()

        return "\\n".join(table_text)

    def _extract_docx_metadata(self, doc: Any) -> Dict[str, Any]:
        """
        提取DOCX特定的元数据

        Args:
            doc: Word文档对象

        Returns:
            Dict[str, Any]: DOCX元数据
        """
        metadata: Dict[str, Any] = {}

        try:
            # 文档核心属性
            core_props = doc.core_properties
            metadata.update(
                {
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "keywords": core_props.keywords or "",
                    "comments": core_props.comments or "",
                    "category": core_props.category or "",
                    "created": core_props.created.isoformat()
                    if core_props.created
                    else "",
                    "modified": core_props.modified.isoformat()
                    if core_props.modified
                    else "",
                    "last_modified_by": core_props.last_modified_by or "",
                    "revision": core_props.revision or "",
                    "version": core_props.version or "",
                }
            )

            # 文档结构信息
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            # 统计非空段落
            non_empty_paragraphs = sum(1 for p in doc.paragraphs if p.text.strip())

            metadata.update(
                {
                    "paragraph_count": paragraph_count,
                    "non_empty_paragraph_count": non_empty_paragraphs,
                    "table_count": table_count,
                }
            )

            # 样式信息
            styles = []
            for paragraph in doc.paragraphs:
                if paragraph.style and paragraph.style.name not in styles:
                    styles.append(paragraph.style.name)

            metadata["styles_used"] = styles[:10]  # 限制样式数量
            metadata["style_count"] = len(styles)

        except Exception as e:
            print(f"提取DOCX元数据时出错: {e}")
            traceback.print_exc()
            metadata["metadata_extraction_error"] = str(e)

        return metadata
